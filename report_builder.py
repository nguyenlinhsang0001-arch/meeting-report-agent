"""
report_builder.py — Sinh 1 BAN BIEN BAN duy nhat tu transcript va xuat Word theo format ARTIUS.
Khong import streamlit -> test doc lap duoc.

DAC DIEM:
  1. Logo / ten cong ty / slogan CO DINH (hang so) -> khong nhan tu UI.
  2. Muc tieu cuoc hop & Thanh phan tham du: AGENT tu doc transcript va rut ra.
  3. So nguoi tham du: tu dem (uu tien so speaker phat hien trong file ghi am).
  4. CHI XUAT 1 FILE: bien ban gon nhung day du cac y chinh (khong con ban tom tat rieng).
"""

import io
import json

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

RED = RGBColor(0xC0, 0x00, 0x00)
GRAY = RGBColor(0x55, 0x55, 0x55)

# ============================================================
# (1) NHAN DIEN CONG TY CO DINH — khong lay tu UI/meta
# ============================================================
COMPANY = "ARTIUS"
TAGLINE = "BEYOND DESIGN AND BUILD"
LOGO_PATH = "logo.jpg"  # dat cung thu muc voi app khi deploy (Streamlit Cloud)


# ============================================================
# GOI CLAUDE -> JSON 1 BAN
# ============================================================
def generate_report(transcript, title, anthropic_key,
                    model="claude-sonnet-4-6"):
    """
    Tra ve 1 dict duy nhat (khong con tach 'full'/'summary').
    Signature da bo tham so `objective` (agent tu rut ra).
    """
    from anthropic import Anthropic  # import lazy

    schema = """{
  "meeting_info": {
    "objective": ["muc tieu cuoc hop - agent TU RUT RA tu noi dung thao luan"],
    "attendees": [
      {"group": "Ten bo phan/nhom (vd: Ban giam doc, BP Du an, Thau phu MEP)",
       "members": ["Mr. A", "Ms. B"]}
    ],
    "attendee_count": 0
  },
  "content_heading": "Ten nhom noi dung (vd: Cac noi dung MEP)",
  "topics": [{"heading": "Ten chu de", "points": ["y chinh 1", "y chinh 2"]}],
  "decisions": ["quyet dinh / dieu da thong nhat trong hop"],
  "action_plan": [{"task": "viec can lam", "owner": "ten/don vi phu trach", "deadline": "ngay cu the hoac 'Chua xac dinh'"}]
}"""

    prompt = f"""Ban la thu ky chuyen nghiep cua cong ty thiet ke va thi cong noi that ARTIUS.
Tu transcript ben duoi (da ghi ro ten + vai tro nguoi noi), tao MOT ban bien ban GON
nhung the hien DAY DU cac y chinh, tra ve DUY NHAT mot JSON hop le (khong markdown) theo schema:

{schema}

Yeu cau:
- "meeting_info.objective": TU SUY LUAN muc tieu cuoc hop tu noi dung thuc te da thao luan
  (KHONG bia, KHONG chep nguyen tieu de). Nhieu muc tieu thi liet ke tung y.
- "meeting_info.attendees": TU RUT RA thanh phan tham du tu transcript, gom nhom theo
  bo phan/don vi va liet ke ten tung nguoi. Chi tinh nguoi THUC SU xuat hien voi tu cach
  nguoi noi trong transcript.
- "meeting_info.attendee_count": dem chinh xac TONG SO nguoi rieng biet xuat hien (so nguoi noi khac nhau).
- "content_heading": dat ten phu hop loai hop.
- "topics": gom noi dung thanh cac CHU DE, moi topic co heading ngan + points la cac Y CHINH (suc tich, khong dai dong).
- "decisions": liet ke quyet dinh / dieu da thong nhat. Neu khong co thi de mang rong [].
- "action_plan": viec can lam kem nguoi phu trach va han chot cu the neu co.

Quy tac chung:
- Viet tieng Viet co dau day du, giu nguyen thuat ngu ky thuat (MEP, FCU, CDT, PCCC...).
- TRUNG THUC voi transcript, KHONG bia them.

Tieu de cuoc hop: {title}

--- TRANSCRIPT ---
{transcript}
"""
    client = Anthropic(api_key=anthropic_key)
    msg = client.messages.create(
        model=model,
        max_tokens=8000,  # du rong, tranh JSON bi cat cut
        messages=[
            {"role": "user", "content": prompt},
            {"role": "assistant", "content": "{"},  # prefill: ep Claude tra JSON ngay tu dau
        ],
    )

    # Neu het token -> bao loi ro rang thay vi JSON hong kho hieu
    if msg.stop_reason == "max_tokens":
        raise ValueError(
            "Phan hoi bi cat cut do qua dai (max_tokens). "
            "Hay tang max_tokens hoac rut gon transcript."
        )

    raw = "{" + msg.content[0].text                 # ghep lai ky tu '{' da prefill
    raw = raw.replace("```json", "").replace("```", "").strip()
    start, end = raw.find("{"), raw.rfind("}")
    if start != -1 and end != -1:
        raw = raw[start:end + 1]                     # chi giu phan { ... } o giua

    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        raise ValueError(
            f"Claude tra ve khong dung JSON ({e}). 300 ky tu dau:\n{raw[:300]}"
        )


# ============================================================
# HELPERS DOCX
# ============================================================
def _bottom_border(paragraph, color="000000", sz="8"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("val", "single"), ("sz", sz), ("space", "4"), ("color", color)):
        bottom.set(qn("w:" + k), v)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _setup(doc):
    s = doc.styles["Normal"]
    s.font.name = "Arial"
    s.font.size = Pt(11)


def _letterhead(doc, company=COMPANY, tagline=TAGLINE, logo_path=LOGO_PATH):
    """(1) Letterhead CO DINH; dat o HEADER -> lap lai dau moi trang, can trai."""
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    used_image = False
    if logo_path:
        try:
            p.add_run().add_picture(logo_path, width=Inches(2.0))
            used_image = True
        except Exception:
            used_image = False
    if not used_image:
        r = p.add_run(" ".join(company.upper()))
        r.font.size = Pt(20)
        r.font.color.rgb = RGBColor(0, 0, 0)
        if tagline:
            sub = header.add_paragraph()
            sr = sub.add_run("   ".join(tagline.upper().split()))
            sr.font.size = Pt(7.5)
            sr.font.color.rgb = GRAY
            _bottom_border(sub)
            return
    _bottom_border(p)


def _section(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run(text)
    r.bold = True
    r.font.size = Pt(12)


def _sub(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.add_run(text).bold = True


def _label(doc, label, value):
    p = doc.add_paragraph()
    p.add_run(label + ": ").bold = True
    p.add_run(value)


def _dash(doc):
    """Gach dau dong kieu '-' co thut hang treo (giong template)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.22)
    pf.space_after = Pt(4)
    p.add_run("-\t")
    return p


def _count_attendees(meta, info):
    """(3) Tu dem: uu tien speaker_count tu diarization, roi den so agent dem, cuoi cung dem thanh vien."""
    n = meta.get("speaker_count")
    if not n:
        n = info.get("attendee_count")
    if not n:
        n = sum(len(g.get("members", [])) for g in info.get("attendees", []))
    return n or 0


# ============================================================
# BAN BIEN BAN DUY NHAT (format ARTIUS)
# ============================================================
def build_report_docx(meta, report):
    doc = Document()
    _setup(doc)
    _letterhead(doc)  # (1) co dinh

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(18)
    tr = t.add_run("BIÊN BẢN CUỘC HỌP"); tr.bold = True; tr.font.size = Pt(18)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(meta.get("title", "").upper()); sr.bold = True; sr.font.size = Pt(12)

    info = report.get("meeting_info", {})

    # 1. DIA DIEM & THANH PHAN THAM DU
    _section(doc, "1. ĐỊA ĐIỂM CUỘC HỌP VÀ THÀNH PHẦN THAM DỰ")
    if meta.get("date"):
        _label(doc, "Ngày họp", meta["date"])
    if meta.get("location"):
        _label(doc, "Địa điểm", meta["location"])

    attendees = info.get("attendees", [])
    count = _count_attendees(meta, info)  # (3)
    p = doc.add_paragraph()
    p.add_run(f"Thành phần tham dự ({count} người):").bold = True
    for g in attendees:  # (2) agent rut ra
        line = _dash(doc)
        if g.get("group"):
            line.add_run(g["group"] + ": ").bold = True
        line.add_run(", ".join(g.get("members", [])))
    if not attendees:
        _dash(doc).add_run("(Agent chưa nhận diện được người tham dự từ file ghi âm)")

    # 2. MUC TIEU
    _section(doc, "2. MỤC TIÊU CUỘC HỌP")
    objectives = info.get("objective", [])
    if isinstance(objectives, str):
        objectives = [objectives]
    for o in objectives or ["(Agent chưa xác định được mục tiêu từ nội dung cuộc họp)"]:
        _dash(doc).add_run(o)

    # 3. NOI DUNG CHINH
    _section(doc, "3. NỘI DUNG CUỘC HỌP")
    _sub(doc, "3.1. " + (report.get("content_heading") or "Các nội dung chính") + ":")
    for topic in report.get("topics", []):
        p = _dash(doc)
        p.add_run(topic.get("heading", "") + ": ").bold = True
        p.add_run(" ".join(topic.get("points", [])))
    if not report.get("topics"):
        _dash(doc).add_run("(Không có)")

    # 3.2 QUYET DINH / THONG NHAT (chi hien khi co)
    decisions = report.get("decisions", [])
    if decisions:
        _sub(doc, "3.2. Quyết định / Thống nhất")
        for d in decisions:
            _dash(doc).add_run(str(d))
        plan_no = "3.3."
    else:
        plan_no = "3.2."

    # KE HOACH TRIEN KHAI
    _sub(doc, plan_no + " Kế hoạch triển khai")
    for a in report.get("action_plan", []):
        p = _dash(doc)
        p.add_run(a.get("task", ""))
        if a.get("owner"):
            p.add_run(f" (Phụ trách: {a['owner']})")
        if a.get("deadline"):
            p.add_run("  Hạn: ")
            p.add_run(a["deadline"]).font.color.rgb = RED
    if not report.get("action_plan"):
        _dash(doc).add_run("(Không có)")

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()
