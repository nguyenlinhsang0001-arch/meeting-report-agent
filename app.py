"""
app.py — Web app tao bien ban cuoc hop (chay local, giao dien than thien)

Stack:
  Streamlit    -> giao dien keo-tha file
  AssemblyAI   -> STT + diarization tieng Viet
  Claude API   -> sinh bien ban dang JSON
  python-docx  -> xuat file Word 1 trang co tieu de cong ty (letterhead)

CAI DAT:
    pip install streamlit assemblyai anthropic python-docx

CHAY:
    streamlit run app.py
(Trinh duyet se tu mo http://localhost:8501)
"""

import io
import json
import tempfile

import streamlit as st
import assemblyai as aai
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# CAU HINH — doc key tu st.secrets (KHONG hardcode vao code)
#   - Chay local: tao file .streamlit/secrets.toml
#   - Deploy:     dan noi dung secrets.toml vao muc "Secrets" tren Streamlit Cloud
# ============================================================
ASSEMBLYAI_API_KEY = st.secrets["ASSEMBLYAI_API_KEY"]
ANTHROPIC_API_KEY = st.secrets["ANTHROPIC_API_KEY"]
CLAUDE_MODEL = st.secrets.get("CLAUDE_MODEL", "claude-sonnet-4-6")  # doi sang opus neu muon chat luong cao hon

aai.settings.api_key = ASSEMBLYAI_API_KEY


# ============================================================
# 1) STT + DIARIZATION
# ============================================================
def transcribe(audio_bytes: bytes, suffix: str, speakers_expected: int):
    """Tra ve list utterances, moi cai co .speaker (A,B..) va .text"""
    # Streamlit cho bytes, AssemblyAI can duong dan file -> ghi tam
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(audio_bytes)
        tmp_path = tmp.name

    config = aai.TranscriptionConfig(
        language_code="vi",
        speaker_labels=True,
        speakers_expected=speakers_expected,
    )
    transcript = aai.Transcriber().transcribe(tmp_path, config)
    if transcript.status == aai.TranscriptStatus.error:
        raise RuntimeError(transcript.error)
    return transcript.utterances


# ============================================================
# 2) SINH BIEN BAN (Claude tra ve JSON co cau truc)
# ============================================================
def generate_minutes_json(labeled_transcript: str, meeting_title: str) -> dict:
    client = Anthropic(api_key=ANTHROPIC_API_KEY)
    prompt = f"""Ban la thu ky cuoc hop chuyen nghiep cua cong ty thiet ke va thi
cong noi that. Tu transcript ben duoi (da ghi ro ten + vai tro nguoi noi), hay
soan bien ban NGAN GON (toi da 1 trang), di thang vao viec.

CEO uu tien doc 2 phan dau tien, nen viet ky 2 phan do va ngan gon cac phan sau.

Tra ve DUY NHAT mot object JSON hop le (khong markdown, khong giai thich) theo schema:
{{
  "summary": "2-3 cau tom tat dieu hanh",
  "risks_approvals": ["rui ro hoac viec can CEO quyet dinh, moi y 1 dong", "..."],
  "progress_budget": ["cap nhat tien do / ngan sach du an, moi y 1 dong", "..."],
  "decisions": ["quyet dinh da thong qua trong cuoc hop", "..."],
  "action_items": [
    {{"owner": "ten - vai tro", "task": "viec can lam", "deadline": "han chot hoac 'Chua xac dinh'"}}
  ]
}}
Neu phan nao khong co thong tin, tra ve mang rong [].

Tieu de cuoc hop: {meeting_title}

--- TRANSCRIPT ---
{labeled_transcript}
"""
    msg = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = msg.content[0].text.strip()
    # Phong khi model bao JSON trong ```json ... ```
    raw = raw.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)


# ============================================================
# 3) XUAT FILE WORD 1 TRANG CO LETTERHEAD
# ============================================================
def _add_bottom_border(paragraph):
    """Them duong ke ngang duoi 1 doan (dung lam letterhead, tranh dung bang)."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2E75B6")
    pbdr.append(bottom)
    pPr.append(pbdr)


def build_docx(data: dict, company: str, meeting_title: str, meeting_date: str) -> bytes:
    doc = Document()

    # Font mac dinh dong nhat
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    # --- Letterhead: ten cong ty + duong ke ---
    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = head.add_run(company.upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    _add_bottom_border(head)

    # --- Tieu de bien ban ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("BIEN BAN CUOC HOP")
    tr.bold = True
    tr.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{meeting_title}  •  {meeting_date}").italic = True

    def heading(text):
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    def bullets(items):
        if not items:
            doc.add_paragraph("(Khong co)")
        for it in items:
            doc.add_paragraph(str(it), style="List Bullet")

    # --- Tom tat ---
    heading("Tom tat dieu hanh")
    doc.add_paragraph(data.get("summary", ""))

    # --- 2 phan CEO uu tien (len dau) ---
    heading("Rui ro & Viec can CEO duyet")
    bullets(data.get("risks_approvals", []))

    heading("Tien do & Ngan sach du an")
    bullets(data.get("progress_budget", []))

    # --- Cac phan con lai ---
    heading("Quyet dinh da thong qua")
    bullets(data.get("decisions", []))

    heading("Action items")
    actions = data.get("action_items", [])
    if actions:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Nguoi phu trach", "Cong viec", "Han chot"
        for a in actions:
            row = table.add_row().cells
            row[0].text = str(a.get("owner", ""))
            row[1].text = str(a.get("task", ""))
            row[2].text = str(a.get("deadline", ""))
    else:
        doc.add_paragraph("(Khong co)")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# GIAO DIEN STREAMLIT
# ============================================================
st.set_page_config(page_title="Bien ban cuoc hop", page_icon="📝")
st.title("📝 Tao bien ban cuoc hop")

# Thong tin chung (cho letterhead)
with st.sidebar:
    st.header("Thong tin cuoc hop")
    company = st.text_input("Ten cong ty", "Cong ty Noi that ABC")
    meeting_title = st.text_input("Tieu de cuoc hop", "Hop giao ban tuan")
    meeting_date = st.text_input("Ngay hop", "")
    n_speakers = st.number_input("So nguoi du kien", 1, 15, 4)

# Buoc 1: upload + transcribe
uploaded = st.file_uploader(
    "Keo-tha file ghi am vao day",
    type=["mp3", "wav", "m4a", "mp4", "aac"],
)

if uploaded and st.button("1️⃣ Tach loi & nhan dien nguoi noi", type="primary"):
    with st.spinner("Dang xu ly (co the mat vai phut)..."):
        suffix = "." + uploaded.name.split(".")[-1]
        utts = transcribe(uploaded.getvalue(), suffix, int(n_speakers))
        # Luu vao session de song qua cac lan rerun
        st.session_state["utts"] = [
            {"speaker": u.speaker, "text": u.text} for u in utts
        ]
        st.success("Xong! Cuon xuong de gan ten nguoi noi.")

# Buoc 2: gan ten/vai tro cho tung speaker
if "utts" in st.session_state:
    utts = st.session_state["utts"]
    speakers = sorted({u["speaker"] for u in utts})

    st.subheader("2️⃣ Gan ten & vai tro")
    st.caption("Nghe lai / doc cau mau roi dien ten. Vi du: 'Anh Nam - TP Thiet ke'")

    mapping = {}
    for spk in speakers:
        sample = next(u["text"] for u in utts if u["speaker"] == spk)
        st.markdown(f"**Speaker {spk}** — _mau:_ \"{sample[:80]}...\"")
        mapping[spk] = st.text_input(f"Ten cho Speaker {spk}", spk, key=f"name_{spk}")

    # Buoc 3: sinh bien ban + tai Word
    st.subheader("3️⃣ Tao bien ban")
    if st.button("📄 Tao bien ban Word", type="primary"):
        labeled = "\n".join(f"[{mapping[u['speaker']]}] {u['text']}" for u in utts)
        with st.spinner("Claude dang viet bien ban..."):
            try:
                data = generate_minutes_json(labeled, meeting_title)
            except json.JSONDecodeError:
                st.error("Claude tra ve khong dung JSON. Thu lai hoac doi sang model opus.")
                st.stop()

        # Xem truoc nhanh tren man hinh
        st.markdown("**Xem truoc:**")
        st.write(data.get("summary", ""))

        docx_bytes = build_docx(data, company, meeting_title, meeting_date)
        st.download_button(
            "⬇️ Tai bien ban (.docx)",
            data=docx_bytes,
            file_name="bien_ban.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        # Luu them transcript tho de doi chieu khi can
        st.download_button(
            "⬇️ Tai transcript tho (.txt)",
            data=labeled.encode("utf-8"),
            file_name="transcript.txt",
            mime="text/plain",
        )
